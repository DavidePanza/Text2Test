import streamlit as st
import json
from app.utils import *
from app.main_IO import *
from app.pages.utils_chapter.page1_utils import *
from app.pages.utils_chapter.select_toc import *
from app.backend.raw_text_processing import *
from app.backend.get_requests import extract_chapters_from_toc, generate_questions_from_chapter, generate_questions_from_chapter_edgecase
from app.backend.text_processing import chapters_chunking
from app.backend.chunks_processing import get_chapter_context



st.title("Generate Questions from a Chapter")
st.write("Here, you can generate questions based on a specific chapter.")

# if st.session_state.get("uploaded_pdf_bytes") is None:
#     st.warning("Please upload a PDF file to generate questions from a chapter.")
# else:
#     pass

level = st.selectbox("Logging level", ["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"])
logging.getLogger().setLevel(getattr(logging, level))

# Display the PDF preview and page range selector
show_pdf_preview()
display_scrollable_pages()

# UI and Interaction
set_clicked, start_page, end_page = page_range_selector_ui()

if set_clicked:
    updated = handle_page_range_submission(start_page, end_page)
    st.session_state["page_range_updated"] = updated

if st.session_state.get("page_range_updated", False):
    extract_content_if_needed()
    st.session_state["page_range_updated"] = False

# Gate rest of app
if st.session_state.get("page_range_set", False):
    # Render rest of app
    ...
else:
    st.info("Please set a valid page range to continue.")

# # Page range selector for TOC extraction
# page_range_selector()

# if st.session_state.get("page_range_set", False):
#     if st.session_state.toc is not None and st.session_state.chapters_dict is None:
#         with st.spinner("Extracting chapters from TOC..."):
#             extract_chapters_from_toc(st.session_state.toc)
#             st.success("Chapters extracted successfully.")

#         extract_chapters(st.session_state['chapters_dict'], st.session_state['pages_data_infos'])
#         try:
#             debug_log(f"Number of chapters extracted: {len(st.session_state['chapters_extracted'])}")
#             debug_log(f"First chapter content preview:\n{st.session_state['chapters_extracted'][0]['content'][:1000]}")
#         except KeyError:
#             debug_log("'chapters_extracted' not found in session state.")
#         except IndexError:
#             debug_log("'chapters_extracted' list is empty.")

# if st.session_state['toc_page_range'] is not None:
#     debug_log(f"Selected page range: {st.session_state['toc_page_range'][0] + 1} to {st.session_state['toc_page_range'][1] + 1}")

#     extract_toc(st.session_state['toc_page_range'])
#     debug_log(f"Table of Contents (TOC): {st.session_state['toc'][:200]}...")

# if st.session_state['toc'] is not None and st.session_state['chapters_dict'] is None:
#     with st.spinner("Extracting chapters from TOC..."):
#         extract_chapters_from_toc(st.session_state['toc'])
#         st.success("Chapters extracted successfully.")

try:
    debug_log(f"Chapters dictionary: {st.session_state['chapters_dict'][0]}")
    debug_log(f"Pages data infos: {st.session_state['pages_data_infos'][0]}")
except:
    pass

extract_chapters(st.session_state['chapters_dict'], st.session_state['pages_data_infos'])

try:
    debug_log(f"Number of chapters extracted: {len(st.session_state['chapters_extracted'])}")
    debug_log(f"First chapter content preview:\n{st.session_state['chapters_extracted'][0]['content'][:1000]}")
except KeyError:
    debug_log("'chapters_extracted' not found in session state.")
except IndexError:
    debug_log("'chapters_extracted' list is empty.")


chapters_chunking(st.session_state['chapters_extracted'])

try:
    debug_log(f"Number of chapters chunked: {len(st.session_state['chapters_chunked'])}")
except:
    debug_log("'chapters_chunked' not found in session state.")


# select chapter
if st.session_state.get('chapters_dict') is not None:
    chapters = st.session_state['chapters_dict']

    if len(chapters) > 0 and isinstance(chapters[0], dict) and 'chapter_title' in chapters[0]:
        chapter_titles = ["Chapter " + str(idx) + ": " + ch['chapter_title'] for idx, ch in enumerate(chapters, start=1)]

    options = st.multiselect(  # use selectobx columns here https://docs.streamlit.io/develop/api-reference/data/st.column_config/st.column_config.selectboxcolumn
        "Select a Chapter:",
        [chapter_titles][0],
        max_selections=1,
        accept_new_options=False,
    )

    st.session_state.selected_chapter_title = options[0]
    st.write(f"Selected chapter: {options[0] if options else 'None'}")

# Get the index of the selected title
if options:
    st.session_state['selected_chapter_idx'] = chapter_titles.index(options[0])
    st.write(f"Selected index: {st.session_state['selected_chapter_idx']}")


num_questions = st.number_input(
    "Number of questions to generate (max 10)",
    min_value=1,
    max_value=5,
    value=None,
    step=1
)

# Optionally store it in session_state
st.session_state['num_questions'] = num_questions

if st.session_state['num_questions'] is not None:
    get_chapter_context(st.session_state['chapters_chunked'],
                        st.session_state['selected_chapter_idx'],
                        num_questions)

    debug_log(f"Selected chapter chunks: {len(st.session_state['chapter_selected_chunks'])}")


if st.session_state['chapter_selected_chunks'] is not None and st.button("Generate Questions"):
    # Replace this with your actual question generation function
    with st.spinner("Generating questions..."):
        if len(st.session_state['chapter_selected_chunks']) >= st.session_state['num_questions']:
            st.session_state['raw_output'] = generate_questions_from_chapter(st.session_state['chapter_selected_chunks'], st.session_state['num_questions'])
        else:
            st.write("use edgecase prompt")
            st.session_state['questions_json'] = generate_questions_from_chapter_edgecase(st.session_state['chapter_selected_chunks'], st.session_state['num_questions'])


breaks(2)
st.header("Generated Questions")
breaks(1)

for idx, question_item in enumerate(st.session_state['questions_json']):
    question_text = question_item['question']
    answer_text = question_item['answer']

    col1, col2 = st.columns([0.9, 0.1])
    
    with col1:
        st.html(f"<p style='font-size:20px; margin:0;'>{idx+1}. {question_text}</p>")
        col1_ = st.columns([1, 4])[0]
        with col1_:
            with st.expander("💡 Show Answer"):
                st.write(answer_text)

    with col2:
        selected = st.checkbox("📌", key=f"select_{idx}", value=True)

breaks(2)

debug_log(f"Questions JSON: {st.session_state.get('questions_json', [])}")

col1_download, col2_download, col3_download = st.columns([0.3, 0.3, 0.6])
with col1_download:
    if st.button("Sync Selected Questions to Download"):
        selected_chapter = st.session_state.get('selected_chapter_title')
        if selected_chapter is None:
            st.error("No chapter selected!")
        else:
            if selected_chapter not in st.session_state['questions_to_download']:
                st.session_state['questions_to_download'][selected_chapter] = []

            current_selected = st.session_state['questions_to_download'][selected_chapter]

            for idx, question in enumerate(st.session_state.get('questions_json', [])):
                current_question = {'question': question['question'], 'answer': question['answer']}
                checkbox_key = f"select_{idx}"
                is_selected = st.session_state.get(checkbox_key, False)

                if is_selected and current_question not in current_selected:
                    current_selected.append(current_question)
                elif not is_selected and current_question in current_selected:
                    current_selected.remove(current_question)

            st.success(f"Selected questions synced for chapter '{selected_chapter}'.")

with col2_download:
    if st.button("Clear Selected Questions"):
        st.session_state['questions_to_download'] = {}
        st.success("Cleared all selected questions.")


# Show what's selected (for testing)
if st.session_state.questions_to_download:
    st.markdown("### ✅ Selected Questions")
    for chapter, questions_list in st.session_state.questions_to_download.items():
        st.markdown(f"#### {chapter}")
        for q in questions_list:
            st.write(q['question'])



from docx import Document
from io import BytesIO
from datetime import datetime

def create_docx_from_data(data):
    doc = Document()
    doc.add_heading("Questions", 0)
    doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    for chapter, qas in data.items():
        doc.add_heading(chapter, level=1)
        doc.add_paragraph("")  # Spacing
        for idx, qa in enumerate(qas, 1):
            doc.add_paragraph(f"Q{idx}: {qa['question']}", style='List Number')
            doc.add_paragraph(f"A: {qa['answer']}", style='Normal')
            doc.add_paragraph("")  # Spacing

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


with st.sidebar:
    st.markdown("---")  # Divider
    st.write("")  # Spacing

    docx_file = create_docx_from_data(st.session_state.get('questions_to_download', {}))

    st.download_button(
        label="📄 Download as Word (.docx)",
        data=docx_file,
        file_name="questions.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


# check https://docs.streamlit.io/develop/api-reference/execution-flow/st.form



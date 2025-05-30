import streamlit as st
from docx import Document
from io import BytesIO
from datetime import datetime


def create_docx_from_data(data):
    doc = Document()
    doc.add_heading("Questions", 0)
    doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    for chapter, qas in data.items():
        doc.add_heading(chapter, level=1)
        doc.add_paragraph("")  # Spacing
        for idx, qa in enumerate(qas, 1):
            doc.add_paragraph(f"Q{idx}: {qa['question']}", style='List Number')
            doc.add_paragraph(f"A: {qa['answer']}", style='Normal')
            doc.add_paragraph("")  # Spacing

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
